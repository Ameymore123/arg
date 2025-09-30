from django.shortcuts import render
from django.shortcuts import render, HttpResponse, redirect, get_object_or_404
from report.models import Prompt
from report.models import Docx_file
from django.utils import timezone
from django.core.files import File

import time

from datetime import datetime

from django.contrib.auth import authenticate, login, logout

from django.contrib.auth.decorators import login_required

from django.contrib.auth.forms import UserCreationForm

from django.contrib.auth.models import User
from icrawler.builtin import BaiduImageCrawler, BingImageCrawler, GoogleImageCrawler

import json
import os
import docx
import google.generativeai as genai
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from requests import get
from bs4 import BeautifulSoup
# Create your views here.


def home(request):
    if request.user.is_authenticated:
        file_url = None
        user = request.user
        if request.method == 'POST':
            p_i = request.POST.get('prompt')
            prompt_input = p_i.lower()
            load_dotenv()


            if not request.session.session_key:
                request.session.save()  # Ensure the session key is generated
            session_key = request.session.session_key

            
            # Initialize Gemini API
            genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
            generation_config = {
                "temperature": 1,
                "top_p": 0.95,
                "top_k": 64,
                "max_output_tokens": 30000,
                "response_mime_type": "text/plain",
            }
            model = genai.GenerativeModel(
                model_name="models/gemini-2.5-flash-preview-09-2025",
                generation_config=generation_config,
                system_instruction="You are a chat bot which is used to generate Projects report of huge paragraphs on given topic, your response should be proper and reliable for storing in a word file in proper format of project report. Use Heading 1 for main sections and Heading 2 for subheadings."
            )

            def generate_report(title, user):
                document = Document()

                # Title
                title_heading = document.add_heading(title.capitalize(), level=1)
                title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Content from Gemini API
                sta = time.time()
                content = fetch_content(title)
                ed = time.time()
                tt = (ed-sta)
                print("Total time to fetch content-------------- ", tt)
                if content:

                    st = time.time()
                    paragraphs = process_content(document, content)
                    nd = time.time()
                    tot = (nd-st)
                    print("Total time to process ----------- ",tot)

                    start = time.time()
                    images = fetch_images(title)
                    end = time.time()
                    total = (end-start)
                    print("Total time to fetch images ----------- ",total)
                    #print(images)

                    # Insert images at suitable places within the paragraphs
                    stttt = time.time()
                    insert_images(document, paragraphs, images)
                    edddd = time.time()
                    tttt = (edddd - stttt)
                    print("Total time to insert imgs---------- ", tttt)
                else:
                    document.add_paragraph("No content available for this topic.")

                # Save the document
                temp_dir = os.path.join('tmp', user.username)
                os.makedirs(temp_dir, exist_ok=True)
                
                filepath = os.path.join(temp_dir, f"{user.username}_{p_i}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx")
                document.save(filepath)

                print("All Total time---------- ", int(tttt + total + tot + tt))

                return filepath

            def fetch_content(title):
                try:
                    response = model.generate_content(f"Generate a detailed and professional Micro Project Report on {title} with proper structure, suitable for engineering students. Include sections such as Introduction in detailed, History in detailed, Working Principle in detailed, Methodology in detailed, Classification in detailed, Types in detailed, Applications in detailed, Results in detailed, Advantages in detailed, Disadvantages in , Conclusion, and References.")
                    
                    if response and response.text:
                        return response.text
                    else:
                        return "No content available for this topic."
                except Exception as e:
                    print(f"Error fetching content: {e}")
                return "No content available for this topic."


            def process_content(document, content):
                paragraphs = []
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()

                    if line.startswith("# "):
                        heading = document.add_heading(line[2:], level=1)
                        for run in heading.runs:
                            run.font.size = Pt(18)
                            run.font.name = 'Arial'
                    elif line.startswith("## "):
                        heading = document.add_heading(line[3:], level=2)
                        for run in heading.runs:
                            run.font.size = Pt(16)
                            run.font.name = 'Arial'
                    elif line.startswith("### "):
                        heading = document.add_heading(line[4:], level=3)
                        for run in heading.runs:
                            run.font.size = Pt(14)
                            run.font.name = 'Arial'

                    elif line.startswith("#### "):
                        heading = document.add_heading(line[5:], level=3)
                        for run in heading.runs:
                            run.font.size = Pt(12)
                            run.font.name = 'Arial'

                    elif line.startswith("* "):
                        p = document.add_paragraph(line[2:].replace("*", ""), style='ListBullet')
                        paragraphs.append(p)
                    elif line.startswith("• "):  # Handling bullet points
                        p = document.add_paragraph(style='ListBullet')
                        parts = line[2:].split("**")
                        for i, part in enumerate(parts):
                            run = p.add_run(part.replace("*", ""))
                            if i % 2 == 1:
                                run.bold = True
                            run.font.size = Pt(12)
                            run.font.name = 'Arial'
                        paragraphs.append(p)
                    else:
                        p = document.add_paragraph()
                        parts = line.split("**")
                        for i, part in enumerate(parts):
                            run = p.add_run(part.replace("*", ""))
                            if i % 2 == 1:
                                run.bold = True
                            run.font.size = Pt(12)
                            run.font.name = 'Arial'
                        paragraphs.append(p)
                return paragraphs



            
            def fetch_images(title):
                tmp_dir = 'tmp'
                os.makedirs(tmp_dir, exist_ok=True)  # Ensure the tmp directory exists

                try:
                    save_dir = os.path.join(tmp_dir, title.replace(' ', '_'))
                    os.makedirs(save_dir, exist_ok=True)  # Create a directory for this title

                    # Use BingImageCrawler to fetch images
                    bing_crawler = BingImageCrawler(downloader_threads=8,
                                                    storage={'root_dir': save_dir})
                    bing_crawler.crawl(keyword=title, filters=None, offset=0, max_num=7)

                    # Return the paths of the downloaded images
                    img_paths = [os.path.join(save_dir, img) for img in os.listdir(save_dir) if img.endswith(('jpg', 'jpeg', 'png'))]
                    return img_paths
                except Exception as e:
                    print(f"Error fetching images: {e}")
                    return []






            def insert_images(document, paragraphs, images):
                for i, paragraph in enumerate(paragraphs):
                    if i % 5 == 0 and i // 5 < len(images):
                        img_path = images[i // 5]
                        try:
                            run = paragraph.add_run()
                            run.add_break()
                            run.add_picture(img_path, width=Inches(4.0))
                            
                        except Exception as e:
                            print(f"Error inserting image: {e}")

                        finally:
                            # Delete the image file after it has been inserted
                            if os.path.exists(img_path):
                                os.remove(img_path)


            # Generate the report
            filepath = generate_report(prompt_input, user)
            print(filepath)

            # Save Report to Database
            with open(filepath, 'rb') as f:
                Docx_file.objects.create(file=File(f), user=user)
                print(p_i + ' created successfully ------------------------------------------------------------------------------------------------------------------------------------')
                

                return redirect('download')
                
        
        return render(request, 'home/index.html', {"user":user})

    else:
        return redirect('sign')





def log(request):
    return render (request, 'home/login.html')




def handlelogin(request):
    if request.method == 'POST':

        loginusername = request.POST['loginusername']
        loginpass = request.POST['loginpass']

        user = authenticate(username=loginusername, password=loginpass)

        if user is not None:

            login(request, user)
            return redirect('/')
        else:
            
            print("invalid credentials")
            print(loginusername)
            print(loginpass)
            return redirect('/log')
    return HttpResponse('404 - Not Found')



def handlelogout(request):

    logout(request)
    return redirect("/log")





def about(request):
    return render(request, 'home/about.html')


def ur(request):
    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["pass"]

        # Check if the username already exists
        if User.objects.filter(username=username).exists():
            print(f"User with username {username} already exists.")
            return redirect("log")  # Redirect to login if user exists

        # Create the user and hash the password
        my_user = User.objects.create_user(username=username, password=password)
        my_user.save()

        print(f"User {my_user.username} created successfully.")
        
    
        user = authenticate(username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect("home")



    else:
        return HttpResponse('404 - Not found')


def sign(request):
    return render(request, 'home/signupp.html')




def download(request):
    file_url = None  # Default value
    if request.method == 'POST':
        user = request.user
        thefile = Docx_file.objects.filter(user=user).last()
        if thefile and thefile.file:  # Ensure file exists
            file_url = thefile.file.url
    elif request.method == 'GET':  # Optionally handle GET requests
        user = request.user
        thefile = Docx_file.objects.filter(user=user).last()
        if thefile and thefile.file:
            file_url = thefile.file.url

    if not file_url:
        message = "No file available for download. Please try again."
    else:
        message = None

    return render(request, 'home/download.html', {"file_url": file_url, "message": message})

